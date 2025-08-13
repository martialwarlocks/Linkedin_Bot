import os
import asyncio
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
import json
import hashlib
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
import uvicorn

# AI and ML imports
import openai
from openai import OpenAI
import tiktoken
from sentence_transformers import SentenceTransformer
import numpy as np
import faiss

# Web crawling imports
import trafilatura
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import newspaper
from readability import Document
import feedparser

# Database and storage
from supabase import create_client, Client
import chromadb
from chromadb.config import Settings

# Document processing
import PyPDF2
import io
from docx import Document as DocxDocument
import fitz  # PyMuPDF

# Environment and config
from dotenv import load_dotenv
import redis
from celery import Celery

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="Enhanced LinkedIn Content Creator API",
    description="Advanced AI-powered LinkedIn content creation with web crawling, document processing, and vector search",
    version="2.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize clients
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
supabase: Client = create_client(
    os.getenv("SUPABASE_URL"),
    os.getenv("SUPABASE_KEY")
)

# Initialize ChromaDB for vector storage
chroma_client = chromadb.PersistentClient(path="./chroma_db")
collection = chroma_client.get_or_create_collection(
    name="linkedin_content_data",
    metadata={"hnsw:space": "cosine"}
)

# Initialize sentence transformer for embeddings
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

# Redis for caching
redis_client = redis.Redis(host='localhost', port=6379, db=0, decode_responses=True)

# Celery for background tasks
celery_app = Celery('linkedin_content_creator')
celery_app.conf.update(
    broker_url='redis://localhost:6379/0',
    result_backend='redis://localhost:6379/0'
)

# Pydantic models
class WebCrawlRequest(BaseModel):
    url: str
    depth: int = Field(default=2, ge=1, le=5)
    max_pages: int = Field(default=10, ge=1, le=50)
    include_links: bool = True

class ContentGenerationRequest(BaseModel):
    prompt: str
    creator_key: str = "gary-v"
    research_ids: List[int] = []
    document_ids: List[str] = []
    crawl_ids: List[int] = []
    max_tokens: int = 2000
    temperature: float = 0.7

class DocumentUploadResponse(BaseModel):
    filename: str
    content_type: str
    size: int
    extracted_text: str
    embeddings_count: int
    document_id: str

class WebCrawlResponse(BaseModel):
    crawl_id: int
    url: str
    pages_crawled: int
    total_content_length: int
    status: str
    created_at: datetime

# Enhanced Web Crawler
class EnhancedWebCrawler:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
    
    async def crawl_website(self, url: str, depth: int = 2, max_pages: int = 10) -> Dict[str, Any]:
        """Enhanced web crawler with multiple extraction methods"""
        try:
            crawl_id = int(datetime.now().timestamp())
            crawled_pages = []
            visited_urls = set()
            
            # Start with the main URL
            urls_to_crawl = [(url, 0)]
            
            while urls_to_crawl and len(crawled_pages) < max_pages:
                current_url, current_depth = urls_to_crawl.pop(0)
                
                if current_url in visited_urls or current_depth > depth:
                    continue
                
                visited_urls.add(current_url)
                
                try:
                    # Extract content using multiple methods
                    content = await self._extract_content(current_url)
                    
                    if content and len(content['text']) > 100:  # Minimum content threshold
                        crawled_pages.append({
                            'url': current_url,
                            'title': content['title'],
                            'text': content['text'],
                            'metadata': content['metadata'],
                            'depth': current_depth
                        })
                        
                        # Store in vector database
                        await self._store_content_in_vector_db(content, current_url, crawl_id)
                        
                        # Find more links if within depth limit
                        if current_depth < depth:
                            links = content.get('links', [])
                            for link in links[:5]:  # Limit links per page
                                if link not in visited_urls:
                                    urls_to_crawl.append((link, current_depth + 1))
                
    except Exception as e:
                    logger.error(f"Error crawling {current_url}: {str(e)}")
                    continue
            
            # Store crawl metadata in Supabase
            crawl_data = {
                'crawl_id': crawl_id,
                'url': url,
                'pages_crawled': len(crawled_pages),
                'total_content_length': sum(len(page['text']) for page in crawled_pages),
                'status': 'completed',
                'created_at': datetime.now().isoformat(),
                'metadata': {
                    'depth': depth,
                    'max_pages': max_pages,
                    'extraction_methods': ['trafilatura', 'newspaper', 'readability']
                }
            }
            
            supabase.table('web_crawls').insert(crawl_data).execute()
            
            return {
                'crawl_id': crawl_id,
                'url': url,
                'pages_crawled': len(crawled_pages),
                'total_content_length': crawl_data['total_content_length'],
                'status': 'completed',
                'created_at': datetime.now(),
                'pages': crawled_pages
            }
            
    except Exception as e:
            logger.error(f"Web crawling failed: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Web crawling failed: {str(e)}")
    
    async def _extract_content(self, url: str) -> Dict[str, Any]:
        """Extract content using multiple methods for better results"""
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            html_content = response.text
            
            content = {
                'title': '',
                'text': '',
                'metadata': {},
                'links': []
            }
            
            # Method 1: Trafilatura (best for article extraction)
            try:
                extracted = trafilatura.extract(html_content, include_formatting=True, include_links=True)
                if extracted:
                    content['text'] = extracted
                    content['title'] = trafilatura.extract_metadata(html_content).get('title', '')
            except Exception as e:
                logger.warning(f"Trafilatura extraction failed: {e}")
            
            # Method 2: Newspaper3k (good for news sites)
            if not content['text'] or len(content['text']) < 200:
                try:
                    article = newspaper.Article(url)
                    article.download(input_html=html_content)
                    article.parse()
                    if article.text and len(article.text) > len(content['text']):
                        content['text'] = article.text
                        content['title'] = article.title or content['title']
                except Exception as e:
                    logger.warning(f"Newspaper extraction failed: {e}")
            
            # Method 3: Readability (fallback)
            if not content['text'] or len(content['text']) < 100:
                try:
                    doc = Document(html_content)
                    content['text'] = doc.summary()
                    content['title'] = doc.title() or content['title']
                except Exception as e:
                    logger.warning(f"Readability extraction failed: {e}")
            
            # Extract links
            try:
                soup = BeautifulSoup(html_content, 'html.parser')
                base_url = response.url
                links = []
                for link in soup.find_all('a', href=True):
                    href = link['href']
                    absolute_url = urljoin(base_url, href)
                    if urlparse(absolute_url).netloc == urlparse(base_url).netloc:
                        links.append(absolute_url)
                content['links'] = list(set(links))[:10]  # Limit to 10 unique links
    except Exception as e:
                logger.warning(f"Link extraction failed: {e}")
            
            # Add metadata
            content['metadata'] = {
                'url': url,
                'extraction_method': 'multi_method',
                'content_length': len(content['text']),
                'links_count': len(content['links'])
            }
            
            return content
            
        except Exception as e:
            logger.error(f"Content extraction failed for {url}: {e}")
            return None
    
    async def _store_content_in_vector_db(self, content: Dict[str, Any], url: str, crawl_id: int):
        """Store extracted content in vector database"""
        try:
            if not content['text']:
                return
            
            # Create embeddings
            embeddings = embedding_model.encode(content['text'])
            
            # Store in ChromaDB
            collection.add(
                embeddings=[embeddings.tolist()],
                documents=[content['text']],
                metadatas=[{
                    'url': url,
                    'title': content['title'],
                    'crawl_id': crawl_id,
                    'type': 'web_crawl',
                    'timestamp': datetime.now().isoformat()
                }],
                ids=[f"crawl_{crawl_id}_{hashlib.md5(url.encode()).hexdigest()}"]
            )
            
        except Exception as e:
            logger.error(f"Failed to store content in vector DB: {e}")

# Enhanced Document Processor
class EnhancedDocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md']
    
    async def process_document(self, file: UploadFile) -> DocumentUploadResponse:
        """Process uploaded document and extract content"""
        try:
            content = await file.read()
            filename = file.filename
            file_extension = Path(filename).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise HTTPException(status_code=400, detail=f"Unsupported file format: {file_extension}")
            
            # Extract text based on file type
            extracted_text = await self._extract_text(content, file_extension)
            
            if not extracted_text or len(extracted_text.strip()) < 50:
                raise HTTPException(status_code=400, detail="Document contains insufficient text content")
            
            # Create document ID
            document_id = hashlib.md5(f"{filename}_{len(content)}".encode()).hexdigest()
            
            # Store in Supabase
            doc_data = {
                'document_id': document_id,
                'filename': filename,
                'content_type': file.content_type,
                'size': len(content),
                'extracted_text': extracted_text,
                'created_at': datetime.now().isoformat(),
                'status': 'processed'
            }
            
            supabase.table('documents').insert(doc_data).execute()
            
            # Create embeddings and store in vector DB
            embeddings_count = await self._store_document_embeddings(extracted_text, document_id, filename)
            
            return DocumentUploadResponse(
                filename=filename,
                content_type=file.content_type,
                size=len(content),
                extracted_text=extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
                embeddings_count=embeddings_count,
                document_id=document_id
            )
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            raise HTTPException(status_code=500, detail=f"Document processing failed: {str(e)}")
    
    async def _extract_text(self, content: bytes, file_extension: str) -> str:
        """Extract text from different file formats"""
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(content)
            elif file_extension == '.docx':
                return self._extract_docx_text(content)
            elif file_extension in ['.txt', '.md']:
                return content.decode('utf-8', errors='ignore')
        else:
                raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
            logger.error(f"Text extraction failed: {e}")
        raise

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
    try:
            doc = fitz.open(stream=content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
        raise

    def _extract_docx_text(self, content: bytes) -> str:
    """Extract text from DOCX"""
    try:
            doc = DocxDocument(io.BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
        raise

    async def _store_document_embeddings(self, text: str, document_id: str, filename: str) -> int:
        """Store document embeddings in vector database"""
        try:
            # Split text into chunks
            chunks = self._split_text_into_chunks(text, chunk_size=1000, overlap=200)
            
            embeddings_count = 0
            for i, chunk in enumerate(chunks):
                if len(chunk.strip()) < 50:  # Skip very short chunks
                    continue
                
                # Create embeddings
                embedding = embedding_model.encode(chunk)
                
                # Store in ChromaDB
                collection.add(
                    embeddings=[embedding.tolist()],
                    documents=[chunk],
                    metadatas=[{
                        'document_id': document_id,
                        'filename': filename,
                        'chunk_index': i,
                        'type': 'document',
                        'timestamp': datetime.now().isoformat()
                    }],
                    ids=[f"doc_{document_id}_{i}"]
                )
                embeddings_count += 1
            
            return embeddings_count
            
    except Exception as e:
            logger.error(f"Failed to store document embeddings: {e}")
            return 0
    
    def _split_text_into_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > chunk_size * 0.7:  # Only break if we're not too early
                    chunk = chunk[:break_point + 1]
                    end = start + break_point + 1
            
            chunks.append(chunk.strip())
            start = end - overlap
        
        return chunks

# Enhanced AI Content Generator
class EnhancedAIContentGenerator:
    def __init__(self):
        self.openai_client = openai_client
        self.creator_styles = {
            "gary-v": {
                "name": "Gary Vaynerchuk",
                "tone": "Direct, passionate, no-nonsense",
                "structure": "Hook → Personal story → Business insight → Call to action",
                "language": "Casual, uses 'you guys', lots of energy",
                "hooks": "Look..., Here's the thing..., Real talk...",
                "endings": "What do you think?, Let me know in the comments, DM me your thoughts"
            },
            "simon-sinek": {
                "name": "Simon Sinek",
                "tone": "Inspirational, thoughtful, leader-focused",
                "structure": "Question → Story/Example → Leadership lesson → Reflection",
                "language": "Professional but warm, thought-provoking",
                "hooks": "Why is it that..., The best leaders..., I once worked with...",
                "endings": "What would you do?, Leadership is a choice, The choice is yours"
            },
            "seth-godin": {
                "name": "Seth Godin",
                "tone": "Wise, concise, marketing-focused",
                "structure": "Insight → Brief explanation → Broader implication",
                "language": "Concise, profound, marketing terminology",
                "hooks": "The thing is..., Here's what I learned..., Marketing is...",
                "endings": "Worth considering., Just saying., Think about it."
            }
        }
    
    async def generate_contextual_content(self, request: ContentGenerationRequest) -> Dict[str, Any]:
        """Generate contextual LinkedIn content using AI with research and document context"""
        try:
            # Gather context from multiple sources
            context_data = await self._gather_context(request)
            
            # Create enhanced prompt
            enhanced_prompt = self._create_enhanced_prompt(request.prompt, request.creator_key, context_data)
            
            # Generate content using OpenAI
            response = await self._generate_with_openai(enhanced_prompt, request.max_tokens, request.temperature)
            
            # Parse and structure the response
            structured_content = self._parse_ai_response(response)
            
            # Store generation metadata
            await self._store_generation_metadata(request, context_data, structured_content)
            
            return structured_content
            
    except Exception as e:
            logger.error(f"Content generation failed: {e}")
            raise HTTPException(status_code=500, detail=f"Content generation failed: {str(e)}")
    
    async def _gather_context(self, request: ContentGenerationRequest) -> Dict[str, Any]:
        """Gather context from research, documents, and crawled content"""
        context_data = {
            'research': [],
            'documents': [],
            'crawled_content': [],
            'total_context_length': 0
        }
        
        try:
            # Get research data
            if request.research_ids:
                research_data = supabase.table('research').select('*').in_('id', request.research_ids).execute()
                context_data['research'] = research_data.data or []
            
            # Get document data
            if request.document_ids:
                doc_data = supabase.table('documents').select('*').in_('document_id', request.document_ids).execute()
                context_data['documents'] = doc_data.data or []
            
            # Get crawled content
            if request.crawl_ids:
                crawl_data = supabase.table('web_crawls').select('*').in_('crawl_id', request.crawl_ids).execute()
                context_data['crawled_content'] = crawl_data.data or []
            
            # Get relevant content from vector database
            vector_results = await self._search_vector_db(request.prompt)
            context_data['vector_results'] = vector_results
            
            # Calculate total context length
            total_length = 0
            for research in context_data['research']:
                total_length += len(str(research.get('findings', '')))
            for doc in context_data['documents']:
                total_length += len(str(doc.get('extracted_text', '')))
            for crawl in context_data['crawled_content']:
                total_length += len(str(crawl.get('content', '')))
            
            context_data['total_context_length'] = total_length
            
            return context_data
            
        except Exception as e:
            logger.error(f"Context gathering failed: {e}")
            return context_data
    
    async def _search_vector_db(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Search vector database for relevant content"""
        try:
            # Create query embedding
            query_embedding = embedding_model.encode(query)
            
            # Search in ChromaDB
            results = collection.query(
                query_embeddings=[query_embedding.tolist()],
                n_results=limit,
                include=['documents', 'metadatas', 'distances']
            )
            
            return [
                {
                    'content': doc,
                    'metadata': meta,
                    'relevance_score': 1 - distance
                }
                for doc, meta, distance in zip(
                    results['documents'][0],
                    results['metadatas'][0],
                    results['distances'][0]
                )
            ]
            
        except Exception as e:
            logger.error(f"Vector search failed: {e}")
            return []
    
    def _create_enhanced_prompt(self, prompt: str, creator_key: str, context_data: Dict[str, Any]) -> str:
        """Create an enhanced prompt with context and creator style"""
        creator_style = self.creator_styles.get(creator_key, self.creator_styles["gary-v"])
        
        # Build context string
        context_parts = []
        
        if context_data['research']:
            research_context = "Research Context:\n" + "\n".join([
                f"- {r.get('topic', 'Unknown')}: {r.get('findings', '')}"
                for r in context_data['research']
            ])
            context_parts.append(research_context)
        
        if context_data['vector_results']:
            vector_context = "Relevant Content:\n" + "\n".join([
                f"- {result['content'][:200]}..."
                for result in context_data['vector_results'][:3]
            ])
            context_parts.append(vector_context)
        
        if context_data['documents']:
            doc_context = "Document Context:\n" + "\n".join([
                f"- {doc.get('filename', 'Unknown')}: {doc.get('extracted_text', '')[:200]}..."
                for doc in context_data['documents']
            ])
            context_parts.append(doc_context)
        
        context_string = "\n\n".join(context_parts)
        
        # Create the enhanced prompt
        enhanced_prompt = f"""
You are an expert LinkedIn content creator specializing in {creator_style['name']}'s style.

Creator Style:
- Tone: {creator_style['tone']}
- Structure: {creator_style['structure']}
- Language: {creator_style['language']}
- Hooks: {creator_style['hooks']}
- Endings: {creator_style['endings']}

Context Information:
{context_string}

User Request: {prompt}

Please create:
1. A LinkedIn post in {creator_style['name']}'s style (2-3 paragraphs)
2. A LinkedIn reel transcript (30-60 seconds)
3. Key talking points for the content
4. Relevant hashtags

Format your response as JSON:
{{
    "linkedin_post": "the post content",
    "linkedin_reel_transcript": "the transcript",
    "talking_points": ["point1", "point2", "point3"],
    "hashtags": ["#hashtag1", "#hashtag2", "#hashtag3"],
    "style_notes": "notes about how the style was applied"
}}
"""
        
        return enhanced_prompt
    
    async def _generate_with_openai(self, prompt: str, max_tokens: int, temperature: float) -> str:
        """Generate content using OpenAI API"""
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert LinkedIn content creator and social media strategist."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_tokens,
                temperature=temperature
            )
            
            return response.choices[0].message.content
            
    except Exception as e:
            logger.error(f"OpenAI API call failed: {e}")
            raise
    
    def _parse_ai_response(self, response: str) -> Dict[str, Any]:
        """Parse AI response and extract structured content"""
        try:
            # Try to parse as JSON
            if response.strip().startswith('{'):
                return json.loads(response)
            
            # Fallback parsing
            lines = response.split('\n')
            content = {
                'linkedin_post': '',
                'linkedin_reel_transcript': '',
                'talking_points': [],
                'hashtags': [],
                'style_notes': ''
            }
            
            current_section = None
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if 'linkedin post' in line.lower():
                    current_section = 'linkedin_post'
                elif 'reel transcript' in line.lower():
                    current_section = 'linkedin_reel_transcript'
                elif 'talking points' in line.lower():
                    current_section = 'talking_points'
                elif 'hashtags' in line.lower():
                    current_section = 'hashtags'
                elif current_section:
                    if current_section in ['linkedin_post', 'linkedin_reel_transcript']:
                        content[current_section] += line + '\n'
                    elif current_section == 'talking_points':
                        if line.startswith('-') or line.startswith('•'):
                            content['talking_points'].append(line[1:].strip())
                    elif current_section == 'hashtags':
                        if '#' in line:
                            hashtags = [tag.strip() for tag in line.split() if tag.startswith('#')]
                            content['hashtags'].extend(hashtags)
            
            return content
            
    except Exception as e:
            logger.error(f"Failed to parse AI response: {e}")
            return {
                'linkedin_post': response,
                'linkedin_reel_transcript': 'Transcript generation failed',
                'talking_points': [],
                'hashtags': [],
                'style_notes': 'Response parsing failed'
            }
    
    async def _store_generation_metadata(self, request: ContentGenerationRequest, context_data: Dict[str, Any], content: Dict[str, Any]):
        """Store generation metadata for analytics"""
        try:
            metadata = {
                'prompt': request.prompt,
                'creator_key': request.creator_key,
                'research_ids': request.research_ids,
                'document_ids': request.document_ids,
                'crawl_ids': request.crawl_ids,
                'context_length': context_data['total_context_length'],
                'generated_at': datetime.now().isoformat(),
                'content_length': len(str(content))
            }
            
            supabase.table('content_generations').insert(metadata).execute()
            
    except Exception as e:
            logger.error(f"Failed to store generation metadata: {e}")

# Initialize components
web_crawler = EnhancedWebCrawler()
document_processor = EnhancedDocumentProcessor()
ai_generator = EnhancedAIContentGenerator()

# API Endpoints
@app.get("/")
async def health_check():
    """Health check endpoint"""
        return {
        "status": "healthy",
        "message": "Enhanced LinkedIn Content Creator API is online and ready",
        "timestamp": datetime.now().isoformat(),
        "version": "2.0.0",
        "features": [
            "Enhanced web crawling with multiple extraction methods",
            "Vector database for semantic search",
            "Document processing (PDF, DOCX, TXT)",
            "Contextual AI content generation",
            "Research and document integration"
        ]
    }

@app.post("/linkedin/crawl")
async def crawl_website(request: WebCrawlRequest):
    """Enhanced web crawling endpoint"""
    try:
        result = await web_crawler.crawl_website(
            url=request.url,
            depth=request.depth,
            max_pages=request.max_pages
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/linkedin/upload-document")
async def upload_document(file: UploadFile = File(...)):
    """Upload and process document"""
    try:
        result = await document_processor.process_document(file)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/linkedin/generate")
async def generate_content(request: ContentGenerationRequest):
    """Generate contextual LinkedIn content"""
    try:
        result = await ai_generator.generate_contextual_content(request)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/linkedin/research")
async def get_research():
    """Get all research data"""
    try:
        result = supabase.table('research').select('*').order('created_at', desc=True).execute()
        return result.data or []
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/linkedin/documents")
async def get_documents():
    """Get all processed documents"""
    try:
        result = supabase.table('documents').select('*').order('created_at', desc=True).execute()
        return result.data or []
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/linkedin/crawls")
async def get_crawls():
    """Get all web crawls"""
    try:
        result = supabase.table('web_crawls').select('*').order('created_at', desc=True).execute()
        return result.data or []
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/linkedin/creators")
async def get_creators():
    """Get all creator styles"""
    try:
        result = supabase.table('creators').select('*').execute()
        return result.data or []
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/linkedin/search")
async def semantic_search(query: str = Form(...), limit: int = Form(5)):
    """Semantic search in vector database"""
    try:
        results = await ai_generator._search_vector_db(query, limit)
        return {"query": query, "results": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000) 